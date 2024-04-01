import requests
from bs4 import BeautifulSoup
import cssutils
from docx import Document
import re

# Assuming this is your selected element
selected_element_selector = 'article > p'

# URL of the news article
url = 'https://betches.com/jackson-hole-where-to-stay-what-to-do/'

# Fetch the content from the URL
headers = {'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36'}
response = requests.get(url, headers=headers)

html_content = response.text

# Parse the HTML content
soup = BeautifulSoup(html_content, 'html.parser')


article_text = soup.find('body').get_text(separator=' ', strip=True)

# Find all linked stylesheets
stylesheets = [link['href'] for link in soup.find_all('link', rel='stylesheet')]

for css_url in stylesheets:
    # You might need to adjust the URL based on how it's defined in the HTML
    css_response = requests.get(css_url)
    css_content = css_response.text

    # Here you would parse css_content to find rules applicable to your selected_element_selector
    # This step is complex and requires implementing or using a CSS parser

    print("CSS Content:", css_content[:100])  # Just printing the first 100 characters for illustration


# Create a new Word document
doc = Document()

# Iterate over elements in the HTML
for element in soup.body:
    if element.name == 'h1':
        # Example: Add a heading
        run = doc.add_heading(level=1).add_run(element.text)
        # Example of applying a simple style (color)
        if "color" in element.attrs.get('style', ''):
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
    elif element.name == 'p':
        # Add a paragraph
        doc.add_paragraph(element.text)
    # Add more conditions for other elements and styles as needed

# Save the document
doc.save('output.docx')


def parse_css(css_string):
    """
    Parses a CSS string and returns a dictionary of selectors and their styles.
    """
    # Matches a basic CSS rule: selector { styles }
    css_rule_pattern = re.compile(r'([^{]+)\{([^}]+)\}', re.MULTILINE)
    rules = {}
    for selector, styles in css_rule_pattern.findall(css_string):
        # Normalize
        selector = selector.strip().lower()
        styles = styles.strip()
        rules[selector] = styles
    return rules

def element_matches_selector(element, selector):
    """
    Checks if a BeautifulSoup element matches a simplified CSS selector.
    Supports type, class, and ID selectors.
    """
    # Element type selector
    if element.name == selector:
        return True
    # Class selector
    if selector.startswith('.'):
        return selector[1:] in element.get('class', [])
    # ID selector
    if selector.startswith('#'):
        return element.get('id') == selector[1:]
    return False

css_string = soup.find('style').string
css_rules = parse_css(css_string)

# Find all paragraphs and check their applicable styles
for p in soup.find_all('p'):
    for selector, styles in css_rules.items():
        if element_matches_selector(p, selector):
            print(f"Selector '{selector}' matches. Styles: {styles}")
