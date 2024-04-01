import requests
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import cssutils
from docx.shared import Pt, Inches, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def fetch_and_apply_header_styles(document, heading, css_string):
    style_sheet = cssutils.parseString(css_string)
    header_styles = None
    for rule in style_sheet:
        if rule.type == rule.STYLE_RULE:
            if 'h1' in rule.selectorText:  # Example: looking for 'h1' styles
                header_styles = rule.style
                break

    if header_styles:
        # Assuming we have found our header style rules
        font_size = header_styles.getPropertyValue('font-size')
        font_family = header_styles.getPropertyValue('font-family')
        # More properties can be fetched as needed

        # Now apply these styles to the document
        heading_paragraph = document.add_heading(level=1)
        run = heading_paragraph.add_run(heading)

        # Convert CSS font size to pt and set it
        if font_size.endswith('px'):
            # Simple conversion, 1px = 0.75pt. This may need adjustments.
            pt_size = float(font_size.replace('px', '')) * 0.75
            run.font.size = Pt(pt_size)

        # Set font family
        run.font.name = font_family

def fetch_and_resize_image(url, target_height):
    print(f"Fetching image: {url}")
    if url.startswith('//'):
        url = 'https:' + url
    response = requests.get(url)
    if response.status_code == 200:
        try:
            image = Image.open(io.BytesIO(response.content))
            original_width, original_height = image.size
            aspect_ratio = original_width / original_height
            new_width = int(target_height * aspect_ratio)
            
            resized_image = image.resize((new_width, target_height), Image.Resampling.LANCZOS)
            print(f"Image resized to: {new_width}x{target_height}, maintaining aspect ratio.")
            return resized_image
        except Image.UnidentifiedImageError:
            print("Failed to identify image. It might not be a valid image file.")
            return None

def insert_image_to_document(document, image):
    img_io = io.BytesIO()
    if image.mode == 'P' or image.mode == 'RGBA':
        image = image.convert('RGB')
    image.save(img_io, format='JPEG')
    img_io.seek(0)
    last_paragraph = document.add_paragraph()
    run = last_paragraph.add_run()
    run.add_picture(img_io, height=Inches(2))
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("Image inserted and centered in document.")

def insert_styled_text_to_document(document, content, keyword=None, styles_dict=None):
    # Keyword filtering for paragraphs
    if content.name == 'p' and keyword:
        # Convert the paragraph and keyword to lowercase for case-insensitive search
        content_text_lower = content.get_text().lower()
        keyword_lower = keyword.lower()
        
        # Skip the paragraph if the keyword is not found
        if keyword_lower not in content_text_lower:
            return
    
    # Handling lists with keyword filtering
    if content.name == 'ul' or content.name == 'ol':
        # Pass the keyword argument to ensure lists are filtered by the keyword
        insert_list_to_document(document, content, keyword)
    
    elif content.name == 'img':
        # Images are skipped here as they are handled elsewhere
        pass
    
    else:
        p = document.add_paragraph()
        
        # Determine and apply CSS style for the content, if applicable
        css_style = None
        if content.get('style'):
            css_style = content['style']
        elif content.get('class') and styles_dict:
            class_styles = [styles_dict[cls] for cls in content['class'] if cls in styles_dict]
            css_style = ' '.join(class_styles)
        
        if css_style:
            spacing = parse_css_spacing(css_style)
            if spacing.get('line_spacing'):
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = spacing['line_spacing']
            if spacing.get('space_before'):
                p.paragraph_format.space_before = spacing['space_before']
            if spacing.get('space_after'):
                p.paragraph_format.space_after = spacing['space_after']
        
        # Add content to the document, applying any necessary text styles
        for element in content.contents:
            if isinstance(element, NavigableString) and element.parent.name != 'img':
                p.add_run(str(element))
            elif element.name and element.name != 'img':
                run_text = element.get_text() if element.text else str(element)
                run = p.add_run(run_text)
                apply_text_styles(run, element)

def insert_list_to_document(document, list_tag, keyword=None):  # Accept keyword argument
    for item in list_tag.find_all('li', recursive=False):  # Find direct children 'li' elements only
        item_text = item.text.strip()
        
        # If a keyword is specified, perform a case-insensitive search for it in the item text
        if keyword:
            # Convert both item text and keyword to lowercase for case-insensitive comparison
            if keyword.lower() not in item_text.lower():
                continue  # Skip this item if the keyword is not found
        
        # Add a paragraph with a list style for items that pass the keyword check (or if no keyword is specified)
        p = document.add_paragraph(style='ListBullet')
        p.add_run(item_text)

        # Apply any specific styles to the list item text if needed
        apply_text_styles(p.runs[0], item)

        # Check for nested lists within the list item and handle them
        for nested_list in item.find_all(['ul', 'ol'], recursive=False):
            insert_list_to_document(document, nested_list, keyword)  # Pass keyword argument to handle nested lists

def apply_text_styles(run, element):
    if element.name in ['strong', 'b']:
        run.bold = True
    elif element.name in ['em', 'i']:
        run.italic = True

def create_word_document_with_article_and_images(url, keyword):
    document = Document()
    heading, article_content, css_styles = fetch_article_content_and_styles(url)

    latest_image_url = None
    
    # Insert heading into the document using CSS styles
    if heading and css_styles:
        fetch_and_apply_header_styles(document, heading, css_styles)

    for content in article_content:
        if content.name == 'img':
            if 'data-src' in content.attrs:
                latest_image_url = content['data-src']
            elif 'src' in content.attrs:
                latest_image_url = content['src']
            # Don't reset latest_image_url to None here; wait until a matching paragraph is found
        elif content.name in ['p', 'ul', 'ol']:
            content_text_lower = content.get_text().lower() if content.name == 'p' else ""
            keyword_lower = keyword.lower()
            # Check for keyword in paragraph, and ensure we have an image to insert
            if keyword_lower in content_text_lower and latest_image_url:
                # Insert the image before the paragraph
                resized_image = fetch_and_resize_image(latest_image_url, 400)
                if resized_image:
                    insert_image_to_document(document, resized_image)
                latest_image_url = None  # Reset latest_image_url after inserting the image
            # Whether the paragraph contains the keyword or not, it gets processed here
            insert_styled_text_to_document(document, content, keyword=keyword)  # Pass keyword argument

    document_path = 'article_with_images.docx'
    document.save(document_path)
    print(f"Document saved as {document_path}")


def apply_css_styles_to_header(document, heading, css_styles):
    # Example of parsing CSS styles and applying to header
    style_sheet = cssutils.parseString(css_styles)
    for rule in style_sheet:
        if rule.type == rule.STYLE_RULE and 'h1' in rule.selectorText:
            font_size = rule.style.getPropertyValue('font-size')
            font_family = rule.style.getPropertyValue('font-family')
            # Create header with styles
            heading_paragraph = document.add_heading(heading, level=1)
            run = heading_paragraph.runs[0]
            if font_size.endswith('px'):
                pt_size = Pt(float(font_size[:-2]) * 0.75)  # Rough conversion, adjust as needed
                run.font.size = pt_size
            if font_family:
                run.font.name = font_family
            break
        

def fetch_article_content_and_styles(url):
    print("Fetching article content and styles...")
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers)
    css_styles = ''
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        # Remove noscript tags and mega-menu elements from the soup
        for noscript_tag in soup.find_all('noscript'):
            noscript_tag.decompose()
        for mega_menu in soup.select('.mega-menu-main'):
            mega_menu.decompose()
        for menu_item in soup.find_all('div', class_='menu'):
            menu_item.decompose()
        # Extract the heading
        heading = soup.find('h1').text.strip()  # Assuming the heading is in an <h1> tag
        # Extract CSS from <style> tags
        for style_tag in soup.find_all('style'):
            css_styles += style_tag.text
        # Extract and concatenate CSS from external stylesheets
        for link_tag in soup.find_all('link', {'rel': 'stylesheet'}):
             css_url = link_tag['href']
             css_response = requests.get(css_url)
             if css_response.status_code == 200:
                 css_styles += css_response.text
        article_content = soup.select('p, img, ul, ol')
        # Here we have already removed mega-menu elements from the soup
        return heading, article_content, css_styles
    else:
        print("Failed to fetch the article content and styles.")
        return None, [], ''

def parse_css_spacing(style_string):
    """Parse the CSS style string to get spacing values."""
    style = cssutils.parseStyle(style_string)
    line_height = style.getPropertyValue('line-height')
    margin_bottom = style.getPropertyValue('margin-bottom')
    margin_top = style.getPropertyValue('margin-top')
    
    # Convert CSS spacing values to Word spacing values
    return {
        'line_spacing': convert_to_word_spacing(line_height),
        'space_before': convert_to_word_spacing(margin_top),
        'space_after': convert_to_word_spacing(margin_bottom)
    }

def convert_to_word_spacing(css_value):
    """Convert CSS spacing value to Word points."""
    if css_value.endswith('px'):
        return Pt(float(css_value.replace('px', '')) * 0.75)
    elif css_value.endswith('pt'):
        return Pt(float(css_value.replace('pt', '')))
    elif css_value.endswith('em'):
        return Pt(float(css_value.replace('em', '')) * 12)  # Assuming 1em = 12pt
    # Add more conversions if necessary
    return None

# Example usage
url = 'https://betches.com/jackson-hole-where-to-stay-what-to-do/'
create_word_document_with_article_and_images(url, "Kemo Sabe")
